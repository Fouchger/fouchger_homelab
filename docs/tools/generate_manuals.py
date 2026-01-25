#!/usr/bin/env python3
"""
Fouchger Homelab – Manual Generator

Description:
  Generates the User Manual, Developer Manual, and a short Operational Runbook
  for the fouchger_homelab repo. It also generates PNG diagrams (architecture,
  menu map, and data model) and embeds them into the DOCX outputs.

Outputs:
  - Fouchger_Homelab_User_Manual.docx
  - Fouchger_Homelab_Developer_Manual.docx
  - Fouchger_Homelab_Runbook.docx
  - manual_assets/architecture.png
  - manual_assets/menu_map.png
  - manual_assets/data_model.png

Optional:
  If LibreOffice is installed (soffice), the script will also export PDFs
  alongside the DOCX files.

Usage:
  python3 generate_manuals.py --repo /path/to/fouchger_homelab-main --out /path/to/output

Notes:
  - Safe-by-default: reads the repo and writes docs/diagrams only; does not modify repo files.
  - Designed for Ubuntu/Debian environments.
  - Keep diagrams simple and readable; this is operational documentation.

Maintainer:
  Your team / you

-------------------------------------------------------------------------------
"""

from __future__ import annotations

import argparse
import datetime as _dt
import os
import re
import shutil
import subprocess
import textwrap
from pathlib import Path
from typing import Dict, List, Tuple

import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


# -----------------------------
# General helpers
# -----------------------------

def die(msg: str) -> None:
    raise SystemExit(f"ERROR: {msg}")


def ensure_dir(p: Path) -> None:
    p.mkdir(parents=True, exist_ok=True)


def read_text(p: Path) -> str:
    return p.read_text(errors="ignore")


def run(cmd: List[str]) -> None:
    """Run a command with basic error reporting."""
    proc = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
    if proc.returncode != 0:
        raise RuntimeError(
            f"Command failed ({proc.returncode}): {' '.join(cmd)}\n"
            f"STDOUT:\n{proc.stdout}\nSTDERR:\n{proc.stderr}"
        )


# -----------------------------
# Bash parsing (best effort)
# -----------------------------

def extract_functions(path: Path) -> List[str]:
    """
    Extract Bash function names using a simple regex.

    Notes:
      - This is intentionally lightweight (no full Bash parsing).
      - It should be used as a navigation aid, not a formal API contract.
    """
    txt = read_text(path)
    fn = re.findall(r'^\s*(?:function\s+)?([A-Za-z_][A-Za-z0-9_]*)\s*\(\)\s*\{', txt, flags=re.M)
    return sorted(set(fn))


def parse_ui_menu_blocks(menu_sh_text: str) -> List[Tuple[str, str, str, List[Tuple[str, str]]]]:
    """
    Parse ui_menu blocks from lib/menu.sh.

    Returns:
      List of tuples: (title, prompt, variable, items)
      where items is list of (key, label)

    Notes:
      - This is best-effort. It assumes the ui_menu usage pattern found in the repo.
      - If the menu formatting changes, this may return fewer results.
    """
    blocks = []
    pattern = r'ui_menu\s+"([^"]+)"\s+"([^"]+)"\s+([A-Za-z_][A-Za-z0-9_]*)\s*\\\s*(.*?)\n\n'
    for m in re.finditer(pattern, menu_sh_text, flags=re.S):
        title, prompt, var, body = m.group(1), m.group(2), m.group(3), m.group(4)
        items = re.findall(r'\n?\s*([0-9]+)\s+"([^"]+)"', body)
        blocks.append((title, prompt, var, items))
    return blocks


# -----------------------------
# Diagram generation
# -----------------------------

def diagram_boxes_arrows(
    out_path: Path,
    boxes: Dict[str, Tuple[float, float, float, float, str]],
    arrows: List[Tuple[str, str, str]],
    *,
    figsize: Tuple[float, float] = (11, 6),
    title: str | None = None,
) -> None:
    """
    Create a simple box-and-arrow diagram.

    boxes:
      name -> (x, y, w, h, text) in 0..1 coordinates

    arrows:
      (from_name, to_name, label)

    Notes:
      - Keep it simple; these are docs diagrams, not a full modelling tool.
      - Avoid overloading the diagram; readability matters more than completeness.
    """
    fig, ax = plt.subplots(figsize=figsize)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")

    if title:
        ax.text(0.5, 0.98, title, ha="center", va="top", fontsize=14, fontweight="bold")

    for name, (x, y, w, h, txt) in boxes.items():
        patch = FancyBboxPatch(
            (x, y),
            w,
            h,
            boxstyle="round,pad=0.02,rounding_size=0.02",
            linewidth=1,
        )
        ax.add_patch(patch)
        ax.text(x + w / 2, y + h / 2, txt, ha="center", va="center", fontsize=10, wrap=True)

    for frm, to, label in arrows:
        x1, y1, w1, h1, _ = boxes[frm]
        x2, y2, w2, h2, _ = boxes[to]

        # Prefer vertical arrows when stacked; otherwise connect centres.
        start = (x1 + w1 / 2, y1) if y2 < y1 else (x1 + w1 / 2, y1 + h1)
        end = (x2 + w2 / 2, y2 + h2) if y2 < y1 else (x2 + w2 / 2, y2)

        ax.annotate("", xy=end, xytext=start, arrowprops=dict(arrowstyle="->", linewidth=1))
        if label:
            mx = (start[0] + end[0]) / 2
            my = (start[1] + end[1]) / 2
            ax.text(mx, my, label, ha="center", va="center", fontsize=9)

    fig.tight_layout()
    fig.savefig(out_path, dpi=200)
    plt.close(fig)


def generate_diagrams(assets_dir: Path) -> Dict[str, Path]:
    """
    Generate architecture, data model, and menu map diagrams.

    Notes:
      - Diagram content is a pragmatic representation of the repo structure.
      - Update this function if the application evolves materially.
    """
    ensure_dir(assets_dir)

    # Architecture
    arch_boxes = {
        "user": (0.05, 0.75, 0.25, 0.15, "Operator\n(SSH / local shell)"),
        "entry": (0.37, 0.75, 0.25, 0.15, "Entry point\nbin/homelab"),
        "lib": (0.69, 0.75, 0.26, 0.15, "Library layer\nlib/*.sh"),
        "ui": (0.69, 0.53, 0.26, 0.15, "UI framework\n(dialog wrappers)"),
        "menu": (0.37, 0.53, 0.25, 0.15, "Menu router\nlib/menu.sh"),
        "actions": (0.37, 0.31, 0.25, 0.15, "Actions\nlib/actions.sh"),
        "modules": (0.05, 0.31, 0.25, 0.15, "Feature modules\nscripts/core/*.sh"),
        "state": (0.69, 0.31, 0.26, 0.15, "State + logs\n~/.config/fouchger_homelab"),
    }
    arch_arrows = [
        ("user", "entry", "run"),
        ("entry", "lib", "source"),
        ("lib", "menu", "expose"),
        ("lib", "ui", "use"),
        ("menu", "actions", "dispatch"),
        ("actions", "modules", "invoke"),
        ("modules", "state", "read/write"),
        ("lib", "state", "read/write"),
        ("ui", "state", "tempfiles"),
    ]
    arch_path = assets_dir / "architecture.png"
    diagram_boxes_arrows(arch_path, arch_boxes, arch_arrows, figsize=(11, 6.5), title="System Architecture (High Level)")

    # Data model
    data_boxes = {
        "stateenv": (0.05, 0.70, 0.35, 0.18, "state.env\nFeature flags and host settings\n(~/.config/fouchger_homelab/state.env)"),
        "appenv": (0.05, 0.42, 0.35, 0.18, "app_install_list.env\nSelections + version pins\n(~/.config/fouchger_homelab/app_manager/...)"),
        "markers": (0.05, 0.14, 0.35, 0.18, "Markers\nWhat this tool installed\n(~/.config/.../state/markers)"),
        "logs": (0.55, 0.56, 0.40, 0.18, "Logs\nLayer 1: app-manager.log etc\n(~/.config/.../logs)"),
        "ptlog": (0.55, 0.28, 0.40, 0.18, "Optional Layer 2 capture\nptlog\n(~/.ptlog/current.log)"),
    }
    data_arrows = [
        ("stateenv", "appenv", "feature gates"),
        ("appenv", "markers", "apply changes\ncreates/removes"),
        ("appenv", "logs", "writes"),
        ("stateenv", "logs", "writes"),
        ("stateenv", "ptlog", "enables\nvia FEATURE_SESSION_CAPTURE"),
    ]
    data_path = assets_dir / "data_model.png"
    diagram_boxes_arrows(data_path, data_boxes, data_arrows, figsize=(11, 6), title="Data Model and State Storage")

    # Menu map (simplified)
    menu_boxes = {
        "main": (0.35, 0.78, 0.30, 0.12, "Main menu\nFouchger_Homelab"),
        "git": (0.05, 0.58, 0.26, 0.12, "Git & GitHub\n(dev-auth.sh)"),
        "boot": (0.37, 0.58, 0.26, 0.12, "Bootstrap Dev Server"),
        "infra": (0.69, 0.58, 0.26, 0.12, "Infrastructure"),
        "work": (0.37, 0.38, 0.26, 0.12, "Workflows"),
        "debug": (0.69, 0.38, 0.26, 0.12, "Debug"),
        "appm": (0.37, 0.18, 0.26, 0.12, "Ubuntu App Manager\n(app_manager_menu)"),
    }
    menu_arrows = [
        ("main", "git", "1"),
        ("main", "boot", "2"),
        ("main", "infra", "3"),
        ("main", "work", "4"),
        ("main", "debug", "5"),
        ("boot", "appm", "2"),
    ]
    menu_path = assets_dir / "menu_map.png"
    diagram_boxes_arrows(menu_path, menu_boxes, menu_arrows, figsize=(11, 6), title="Menu Map (Current Wiring)")

    return {"architecture": arch_path, "data_model": data_path, "menu_map": menu_path}


# -----------------------------
# DOCX composition helpers
# -----------------------------

def set_doc_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for h, size in [("Heading 1", 20), ("Heading 2", 16), ("Heading 3", 13)]:
        st = doc.styles[h]
        st.font.name = "Calibri"
        st.font.size = Pt(size)


def add_title(doc: Document, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(28)

    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(12)


def add_footer_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)


# -----------------------------
# Manual generation
# -----------------------------

def build_function_index(repo: Path) -> Dict[str, List[str]]:
    """
    Build a function index for the Developer Manual.

    Notes:
      - Includes *.sh and bin/homelab.
      - Keep this small and practical; long lists are truncated in the doc.
    """
    bash_files = sorted([p for p in repo.rglob("*.sh") if p.is_file()])
    maybe_entry = repo / "bin" / "homelab"
    if maybe_entry.exists():
        bash_files.append(maybe_entry)

    out: Dict[str, List[str]] = {}
    for p in bash_files:
        rel = str(p.relative_to(repo))
        out[rel] = extract_functions(p)
    return out


def create_user_manual(out_dir: Path, diagrams: Dict[str, Path], snapshot_label: str, today: _dt.date) -> Path:
    doc = Document()
    set_doc_styles(doc)
    add_title(doc, "Fouchger Homelab", f"User Manual (v0.1) | {today.strftime('%d %B %Y')}")

    doc.add_paragraph(
        "This manual explains how to install and operate the fouchger_homelab interactive CLI application "
        "on Debian/Ubuntu hosts (including Ubuntu 24.04+ LXC on Proxmox)."
    )

    doc.add_heading("1. What the application does", level=1)
    doc.add_paragraph(
        "fouchger_homelab is an interactive, menu-driven CLI that helps you bootstrap and manage common tooling "
        "for a homelab environment. It provides guided flows for Git and GitHub configuration, a development server "
        "bootstrap path, an Ubuntu App Manager for selecting and applying packages and tools, and a debug area for session capture."
    )

    doc.add_heading("2. Prerequisites", level=1)
    for bullet in [
        "A Debian/Ubuntu-based system. Ubuntu 24.04+ is the primary target for the App Manager flows.",
        "Network access for package installs and any third-party installers you choose to run.",
        "sudo access (or run as root) for installing packages and writing system-level changes.",
        "A terminal that can display dialog-style interfaces (the UI uses dialog under the hood).",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("3. Installation", level=1)
    doc.add_paragraph("There are two common ways to run the tool: directly from a cloned repository, or via the one-liner installer referenced in the repo README.")

    doc.add_heading("3.1 Install from the repository", level=2)
    for step in [
        "Clone the repository to your target host.",
        "From the repo root, run: make menu",
        "If you do not have make installed, install it first (sudo apt-get install -y make) or run the entry point directly: ./bin/homelab",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("3.2 Install via installer script", level=2)
    doc.add_paragraph(
        "The README references a curl-based installer. Treat this as a privileged action and review the script "
        "before running it in production environments."
    )
    doc.add_paragraph('Command: bash -c "$(curl -fsSL https://raw.githubusercontent.com/Fouchger/fouchger_homelab/main/install.sh)"')

    doc.add_heading("4. First-run quick start", level=1)
    for step in [
        "Launch the application (make menu or ./bin/homelab).",
        "Select 'Git & Github Management' to configure your git identity and authenticate GitHub CLI (gh).",
        "Select 'Bootstrap Development Server (admin01)' if you are standing up a dev box and want Code-Server and baseline tooling.",
        "Open 'Debug' to enable session capture if you want full terminal recording (optional).",
        "Use 'Ubuntu App Manager' to pick a profile, adjust selections, and then apply installs.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("5. Navigating the menus", level=1)
    doc.add_paragraph("The current menu map is below. Menu options may be gated by feature flags stored in your state file.")
    doc.add_picture(str(diagrams["menu_map"]), width=Inches(6.8))

    doc.add_heading("6. Main menu options", level=1)
    doc.add_heading("6.1 Git & GitHub Management", level=2)
    doc.add_paragraph("Runs the developer authentication helper to set global git identity (safe-by-default) and authenticate GitHub CLI (gh).")
    doc.add_paragraph("Tip: For non-interactive runs, supply environment variables such as GIT_USER_NAME, GIT_USER_EMAIL, and GITHUB_TOKEN.")

    doc.add_heading("6.2 Bootstrap Development Server", level=2)
    doc.add_paragraph(
        "Provides a guided flow for bootstrapping a dev server. One option downloads and runs a third-party "
        "Code-Server installer. The second option leads into the Ubuntu App Manager to install baseline tooling."
    )
    doc.add_paragraph("Control: the tool will prompt before running third-party scripts.")

    doc.add_heading("6.3 Infrastructure and Workflows", level=2)
    doc.add_paragraph(
        "These menu areas are wired for future capability modules (Proxmox templates, MikroTik integration, DNS services, questionnaires). "
        "They are feature-flagged and may display enablement instructions if disabled on the host."
    )

    doc.add_heading("6.4 Debug", level=2)
    doc.add_paragraph(
        "Includes controls for Layer 2 session capture via ptlog. When enabled, the tool will attempt to start ptlog "
        "automatically on next launch and provide status views and log tails."
    )

    doc.add_heading("7. Ubuntu App Manager (step-by-step)", level=1)
    doc.add_paragraph(
        "The App Manager helps you maintain a repeatable set of packages and tools for Ubuntu 24.04+ hosts, "
        "especially LXC containers. It stores selections and version pins in an env file and tracks installed-by-tool items using marker files."
    )

    doc.add_heading("7.1 Apply a profile (replace selections)", level=2)
    for step in [
        "Open: Bootstrap Development Server (admin01) then 'Bootstrap server - Configs and Setup' to enter the App Manager menu.",
        "Choose 'Apply profile (replace selections)'.",
        "Select a profile (for example Basic, Dev, Automation, Platform).",
        "Confirm the action. This overwrites prior selections in app_install_list.env with the profile defaults.",
        "Optionally adjust selections using 'Change selections' before applying.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("7.2 Apply a profile (add to selections)", level=2)
    for step in [
        "Choose 'Apply profile (add to selections)'.",
        "Select a profile. The profile apps will be added to your current selection set.",
        "Review the updated selection list if prompted, then continue.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("7.3 Change selections", level=2)
    for step in [
        "Choose 'Change selections'.",
        "Tick or untick apps using the checklist.",
        "Save and return to the App Manager menu.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("7.4 Edit version pins", level=2)
    for step in [
        "Choose 'Edit version pins'.",
        "Set versions to 'latest' or a specific value (where supported).",
        "Save. Version pins are written to app_install_list.env and used by installers that support pinning.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("7.5 Apply install/uninstall", level=2)
    for step in [
        "Choose 'Apply install/uninstall'.",
        "The tool will compute the delta between selected apps and currently installed-by-tool markers.",
        "Confirm to proceed. Package installs use nala when available (apt-get fallback).",
        "When complete, review the log file if anything failed.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("7.6 Audit which apps are installed", level=2)
    doc.add_paragraph("Choose 'check which apps are installed' to see what the App Manager believes is installed, based on markers and strategy checks.")

    doc.add_heading("8. Where files are stored", level=1)
    doc.add_picture(str(diagrams["data_model"]), width=Inches(6.8))
    doc.add_paragraph("Key paths (defaults):")
    for b in [
        "~/.config/fouchger_homelab/state.env (feature flags, host settings)",
        "~/.config/fouchger_homelab/app_manager/app_install_list.env (selections and version pins)",
        "~/.config/fouchger_homelab/app_manager/app-manager.log (Layer 1 logs)",
        "~/.config/fouchger_homelab/state/markers (installed-by-tool markers)",
        "~/.ptlog/current.log (optional Layer 2 session capture log)",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("9. Troubleshooting", level=1)
    for item in [
        ("The UI does not open", "Ensure dialog is installed and you are in an interactive terminal. If needed, install dialog: sudo apt-get install -y dialog."),
        ("Installs fail due to permissions", "Run as root or ensure sudo is available and your user is in the sudo group."),
        ("GitHub auth fails", "Check gh is installed and your token has appropriate scopes for your workflow. For GHES, ensure GH_HOST is set."),
        ("Session capture does not start", "Install ptlog and enable the feature flag: state_set FEATURE_SESSION_CAPTURE 1, then relaunch the app."),
    ]:
        doc.add_paragraph(f"{item[0]}: {item[1]}", style="List Bullet")

    doc.add_heading("10. Operational guardrails", level=1)
    doc.add_paragraph(
        "The tool aims to be safe-by-default, but it can install packages and run scripts with elevated privileges. "
        "In a corporate environment, treat it like any other automation: review changes, pin versions when stability matters, "
        "and apply in lower environments first."
    )

    add_footer_note(doc, f"Document generated from repository snapshot: {snapshot_label}.")
    out_path = out_dir / "Fouchger_Homelab_User_Manual.docx"
    doc.save(out_path)
    return out_path


def create_developer_manual(out_dir: Path, repo: Path, diagrams: Dict[str, Path], snapshot_label: str, today: _dt.date) -> Path:
    func_index = build_function_index(repo)

    doc = Document()
    set_doc_styles(doc)
    add_title(doc, "Fouchger Homelab", f"Developer Manual (v0.1) | {today.strftime('%d %B %Y')}")

    doc.add_paragraph(
        "This manual describes how the application is structured, how the codebase works, and how to extend it safely. "
        "It also includes a practical code review and recommended improvements."
    )

    doc.add_heading("1. Architectural overview", level=1)
    doc.add_paragraph(
        "fouchger_homelab is a Bash-based, menu-driven CLI. The application follows a layered structure: entry point, "
        "library layer, menu routing, action orchestration, and optional feature modules. State and configuration are "
        "externalised to user-space files for portability."
    )
    doc.add_picture(str(diagrams["architecture"]), width=Inches(6.8))

    doc.add_heading("2. Repository layout", level=1)
    tbl = doc.add_table(rows=1, cols=2)
    hdr = tbl.rows[0].cells
    hdr[0].text = "Path"
    hdr[1].text = "Purpose"
    for p, purpose in [
        ("bin/homelab", "Entrypoint: loads libs and modules, optionally starts session capture, opens main menu."),
        ("lib/*.sh", "Reusable libraries: paths, logging, core helpers, UI wrappers, feature flags, menu routing."),
        ("scripts/core/*.sh", "Operational scripts and feature modules invoked from menus."),
        ("tests/*", "Lightweight tests (currently repo root resolution)."),
        ("install.sh / Makefile", "Installer and developer convenience targets."),
    ]:
        row = tbl.add_row().cells
        row[0].text = p
        row[1].text = purpose

    doc.add_heading("3. Startup sequence", level=1)
    doc.add_paragraph("On launch, bin/homelab anchors REPO_ROOT, loads libraries, loads feature modules, optionally starts Layer 2 capture, then enters main_menu().")
    for step in [
        "Resolve and export REPO_ROOT.",
        "Source lib/modules.sh then call homelab_load_lib (sources lib/paths.sh, logging, core, run, state, common, ui, features, actions, menu).",
        "Call homelab_load_modules (best-effort sources scripts/core/app_manager.sh and other optional modules).",
        "If FEATURE_SESSION_CAPTURE is enabled, attempt to start ptlog and show status.",
        "Call main_menu (lib/menu.sh), which initialises UI and enters a selection loop.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("4. Design principles used in the codebase", level=1)
    for b in [
        "Strict mode (set -Eeuo pipefail) with controlled, explicit defaults.",
        "Best-effort optional modules: missing files do not break core runtime.",
        "UI separation: menu structure in lib/menu.sh; operational orchestration in lib/actions.sh; low-level operations in scripts/*.",
        "State externalisation: persist settings and selections under ~/.config/fouchger_homelab instead of inside the repo.",
        "Marker-based safety: uninstall operations should target only items installed by this tool.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("5. UI framework and menu design", level=1)
    doc.add_paragraph("The UI uses wrapper functions in lib/ui.sh to avoid scattering direct dialog calls through the code. lib/menu.sh focuses on navigation and delegates work to action_* functions (or scripts) to keep responsibilities clean.")
    doc.add_paragraph("Current menu map:")
    doc.add_picture(str(diagrams["menu_map"]), width=Inches(6.8))

    doc.add_heading("6. State, configuration, and data model", level=1)
    doc.add_paragraph("The project uses dotenv-style env files for persistence. This keeps the runtime dependency footprint low and works well in LXC environments.")
    doc.add_picture(str(diagrams["data_model"]), width=Inches(6.8))

    doc.add_heading("7. App Manager internals", level=1)
    doc.add_paragraph(
        "scripts/core/app_manager.sh is a self-contained module that exposes app_manager_menu. "
        "It maintains a catalogue of apps (APP_CATALOGUE) and a set of profiles that map to app keys. "
        "Selections are persisted to app_install_list.env as APP_<KEY>=0/1 along with version variables."
    )

    doc.add_heading("8. API reference (functions and entry points)", level=1)
    doc.add_paragraph(
        "This reference is generated from the current repository snapshot and lists Bash functions discovered in each file. "
        "Use it as a navigation aid rather than a formal interface contract."
    )

    for rel, fns in sorted(func_index.items()):
        doc.add_heading(rel, level=2)
        if not fns:
            doc.add_paragraph("No Bash functions detected.", style="List Bullet")
            continue
        for fn in fns[:60]:
            doc.add_paragraph(fn, style="List Bullet")
        if len(fns) > 60:
            doc.add_paragraph(f"(Truncated: {len(fns)} total functions in this file.)")

    doc.add_heading("9. Setup and configuration", level=1)
    doc.add_heading("9.1 Local development setup (Ubuntu)", level=2)
    for step in [
        "Install prerequisites: git, make, dialog (and optionally gh).",
        "Clone the repo and run make menu to verify the UI starts.",
        "Run make executable (or scripts/core/make-executable.sh) after pulling changes on systems that strip executable bits.",
        "Use scripts/core/bootstrap.sh for minimal bootstrap on fresh hosts.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("10. Tutorials and guides", level=1)
    doc.add_heading("10.1 Add a new main menu item", level=2)
    for step in [
        "Implement a new script or action function (prefer scripts/<area>/... for larger features).",
        "Expose an action_* wrapper in lib/actions.sh if it improves decoupling.",
        "Add a new option in lib/menu.sh main_menu() using ui_menu and dispatch to your action.",
        "If the feature is host-dependent, gate it behind feature_require in lib/features.sh.",
        "Update documentation and add a smoke test if possible.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("11. Code review findings and recommendations", level=1)
    doc.add_heading("11.1 What is working well", level=2)
    for b in [
        "Consistent strict-mode usage across most files improves runtime safety.",
        "Clear separation of menu routing (lib/menu.sh) from operations (lib/actions.sh and scripts).",
        "Repo root resolution and state storage design support installs without a .git directory, which is pragmatic for copied deployments.",
        "Two-layer logging is a strong operational pattern for troubleshooting in homelabs.",
        "App Manager catalogue and profile approach is scalable for repeatable builds.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("11.2 Key risks and improvement opportunities", level=2)
    for b in [
        "Third-party script execution: bootstrap flows may run external curl|bash installers. Consider pinning to a specific commit/tag, logging the exact source URL, and requiring explicit acknowledgement.",
        "Dependency detection: dialog availability is assumed. Consider a preflight that checks for dialog and offers to install it (or falls back to a non-UI mode).",
        "Marker model: decide on one marker location and document it to avoid multi-user surprises.",
        "Idempotency and rollback: repo additions benefit from explicit cleanup paths and backup of repo list files.",
        "Testing: add a lightweight CI check (bash -n, load order, temp HOME env write).",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("12. Code commenting and file header standard", level=1)
    doc.add_paragraph(
        "To keep the codebase consistent, use a standard header block at the top of each script and capture: filename, purpose, usage, "
        "key assumptions, and maintainer. Keep function-level comments focused on why and constraints rather than restating what the code does."
    )
    doc.add_paragraph("Recommended header template:")
    doc.add_paragraph(
        textwrap.dedent(
            """\
            #!/usr/bin/env bash
            # -----------------------------------------------------------------------------
            # Filename: <path>
            # Description: <one line>
            # Usage: <how to run>
            # Notes:
            #   - <key constraints and guardrails>
            # Maintainer: <name/team>
            # -----------------------------------------------------------------------------
            """
        )
    )

    add_footer_note(doc, f"Document generated from repository snapshot: {snapshot_label}.")
    out_path = out_dir / "Fouchger_Homelab_Developer_Manual.docx"
    doc.save(out_path)
    return out_path


def create_runbook(out_dir: Path, snapshot_label: str, today: _dt.date) -> Path:
    doc = Document()
    set_doc_styles(doc)

    doc.add_heading("Fouchger Homelab – Operational Runbook", level=0)
    doc.add_paragraph(f"Quick reference for day-to-day operations | {today.strftime('%d %B %Y')}")

    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph(
        "This runbook provides a concise, operational view of fouchger_homelab for administrators. "
        "It is intended for repeatable day-to-day use, incident response, and first-line troubleshooting."
    )

    doc.add_heading("2. Common daily tasks", level=1)
    for t in [
        "Launch the tool: make menu or ./bin/homelab",
        "Update installed tooling: Ubuntu App Manager → Apply install/uninstall",
        "Add new tools to a host: Ubuntu App Manager → Apply profile (add to selections)",
        "Bootstrap a new dev host: Bootstrap Development Server (admin01)",
        "Verify Git/GitHub auth: Git & GitHub Management",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    doc.add_heading("3. Pre-flight checklist (before changes)", level=1)
    for c in [
        "Confirm you are on the correct host or container",
        "Ensure network connectivity (apt repositories, GitHub, third-party repos)",
        "Check available disk space",
        "Confirm sudo access",
        "Optional: enable session capture for auditability",
    ]:
        doc.add_paragraph(c, style="List Bullet")

    doc.add_heading("4. Where to check when things go wrong", level=1)
    doc.add_paragraph("Primary locations:")
    for p in [
        "~/.config/fouchger_homelab/logs/ (general runtime logs)",
        "~/.config/fouchger_homelab/app_manager/app-manager.log (App Manager actions)",
        "~/.config/fouchger_homelab/state.env (feature flags and settings)",
        "~/.config/fouchger_homelab/state/markers/ (installed-by-tool indicators)",
        "~/.ptlog/current.log (if session capture is enabled)",
    ]:
        doc.add_paragraph(p, style="List Bullet")

    doc.add_heading("5. Common issues and fast recovery", level=1)
    for title, fix in [
        ("UI does not open", "Install dialog and re-run the tool. sudo apt-get install -y dialog"),
        ("Package install failed", "Re-run Apply install/uninstall. Check app-manager.log for the failing package."),
        ("Wrong apps installed", "Review app_install_list.env, adjust selections, then re-apply."),
        ("Tool tries to uninstall something critical", "Stop immediately. Verify marker files before continuing."),
        ("Session capture missing", "Enable FEATURE_SESSION_CAPTURE and restart the tool."),
    ]:
        doc.add_paragraph(f"{title}: {fix}", style="List Bullet")

    doc.add_heading("6. Safe operating guardrails", level=1)
    for g in [
        "Treat this tool as privileged automation; review prompts carefully.",
        "Avoid running bootstrap actions repeatedly on the same host unless intended.",
        "Prefer profiles and version pins for stable environments.",
        "Test changes in a non-production container or VM first.",
    ]:
        doc.add_paragraph(g, style="List Bullet")

    doc.add_heading("7. Escalation and next steps", level=1)
    doc.add_paragraph(
        "If an issue cannot be resolved quickly: stop further changes, collect logs, and review the Developer Manual for deeper diagnostics. "
        "For structural issues, raise a change to the codebase rather than applying manual fixes."
    )

    add_footer_note(doc, f"Document generated from repository snapshot: {snapshot_label}.")
    out_path = out_dir / "Fouchger_Homelab_Runbook.docx"
    doc.save(out_path)
    return out_path


# -----------------------------
# Optional PDF export
# -----------------------------

def export_pdf_with_libreoffice(docx_path: Path, pdf_dir: Path) -> Path | None:
    """
    Convert DOCX to PDF using LibreOffice (soffice) if available.

    Notes:
      - This is the most portable approach in Linux environments.
      - If soffice isn't installed, we skip PDF generation cleanly.
    """
    soffice = shutil.which("soffice")
    if not soffice:
        return None

    ensure_dir(pdf_dir)
    run([soffice, "--headless", "--convert-to", "pdf", "--outdir", str(pdf_dir), str(docx_path)])
    return pdf_dir / (docx_path.stem + ".pdf")


# -----------------------------
# Main
# -----------------------------

def main() -> None:
    parser = argparse.ArgumentParser(description="Generate fouchger_homelab manuals and diagrams.")
    parser.add_argument("--repo", type=Path, default=Path.cwd(), help="Path to repo root (default: current directory)")
    parser.add_argument("--out", type=Path, default=Path.cwd(), help="Output directory for docs and assets")
    parser.add_argument("--snapshot-label", type=str, default=None, help="Label to embed in footers (default: repo name + date)")
    args = parser.parse_args()

    repo = args.repo.resolve()
    out_dir = args.out.resolve()

    if not repo.exists():
        die(f"Repo path does not exist: {repo}")
    if not (repo / "lib").exists() or not (repo / "bin").exists():
        die(f"Repo path does not look like the expected project root (missing lib/ or bin/): {repo}")

    ensure_dir(out_dir)
    assets_dir = out_dir / "manual_assets"
    pdf_dir = out_dir / "manual_out"

    today = _dt.date.today()
    snapshot_label = args.snapshot_label or f"{repo.name} ({today.strftime('%d %B %Y')})"

    diagrams = generate_diagrams(assets_dir)

    user_docx = create_user_manual(out_dir, diagrams, snapshot_label, today)
    dev_docx = create_developer_manual(out_dir, repo, diagrams, snapshot_label, today)
    runbook_docx = create_runbook(out_dir, snapshot_label, today)

    # Optional PDF export
    exported = []
    for docx in [user_docx, dev_docx, runbook_docx]:
        pdf = export_pdf_with_libreoffice(docx, pdf_dir)
        if pdf:
            exported.append(pdf)

    print("Generated DOCX:")
    print(f"  - {user_docx}")
    print(f"  - {dev_docx}")
    print(f"  - {runbook_docx}")

    if exported:
        print("Generated PDF:")
        for p in exported:
            print(f"  - {p}")
    else:
        print("PDF export skipped (LibreOffice 'soffice' not found).")


if __name__ == "__main__":
    main()
